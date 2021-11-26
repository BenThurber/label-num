import re
import os.path
from docx import Document

import tkinter
from tkinter.filedialog import askopenfilename


def get_file_from_dialogue():
    tkinter.Tk().withdraw()
    filename = askopenfilename()
    return filename


def paragraph_replace_text(paragraph, regex, start_num, opening_char, closing_char):
    """Return `paragraph` after replacing all matches for `regex` with `start_num`.

    `regex` is a compiled regular expression prepared with `re.compile(pattern)`
    according to the Python library documentation for the `re` module.
    """
    # --- store how many times the string was replaced ---
    count = 0
    # --- a paragraph may contain more than one match, loop until all are replaced ---
    for match in regex.finditer(paragraph.text):
        match_num = int(match.group(0).lstrip(opening_char).rstrip(closing_char))
        label_num = start_num + match_num
        # --- calculate how much characters must be shifted to fix the match ---
        padding = (len(str(label_num)) - (match.end() -match.start()) ) *count

        # --- when there's a match, we need to modify run.text for each run that
        # --- contains any part of the match-string.
        runs = iter(paragraph.runs)
        start, end = match.start() + padding , match.end() + padding

        # --- Skip over any leading runs that do not contain the match ---
        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start, end = start - run_len, end - run_len

        # --- Match starts somewhere in the current run. Replace match-str prefix
        # --- occurring in this run with entire replacement str.
        run_text = run.text
        run_len = len(run_text)
        run.text = "%s%s%s" % (run_text[:start], str(label_num), run_text[end:])
        end -= run_len  # --- note this is run-len before replacement ---

        # --- Remove any suffix of match word that occurs in following runs. Note that
        # --- such a suffix will always begin at the first character of the run. Also
        # --- note a suffix can span one or more entire following runs.
        for run in runs:  # --- next and remaining runs, uses same iterator ---
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len
        count += 1
    # --- optionally get rid of any "spanned" runs that are now empty. This
    # --- could potentially delete things like inline pictures, so use your judgement.
    # for run in paragraph.runs:
    #     if run.text == "":
    #         r = run._r
    #         r.getparent().remove(r)
    return paragraph


def select_file():
    filetypes = (
        ('text files', '*.txt'),
        ('All files', '*.*')
    )

    filename = fd.askopenfilename(
        title='Open a file',
        initialdir='/',
        filetypes=filetypes)

    showinfo(
        title='Selected File',
        message=filename
    )



def main():
    
    file_path = get_file_from_dialogue()
    
    f = open(file_path, 'rb')
    document = Document(f)
    f.close()
    
    regex = re.compile("<[0-9]+>")
    
    print("------------------------------")
    print("----Label-Number-Generator----")
    print("------------------------------")
    print()
    print("Generated file is created next to origional file.")
    start_num = int(input("Enter the label number to start at: "))
    
    for paragraph in document.paragraphs:
        paragraph = paragraph_replace_text(paragraph, regex, start_num, '<', '>')
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                print(dir(cell))
                for paragraph in cell.paragraphs:
                    paragraph = paragraph_replace_text(paragraph, regex, start_num, '<', '>')
    
    
    
    file_name, file_extension = os.path.splitext(file_path)
    file_path_generated = file_name + "_GENERATED" + file_extension
    if os.path.exists(file_path_generated):
        answer = input('File "{}" alreay exists.  Overwrite? (Y/n): '.format(file_path_generated))
        if answer.upper() != 'Y':
            print("Aborting...")
            return
    
    print("Writing to {}".format(os.path.basename(file_path_generated)))
    document.save(file_name + "_GENERATED" + file_extension)
    print("Done.  Exiting.")



if __name__ == "__main__":
    
    main()

